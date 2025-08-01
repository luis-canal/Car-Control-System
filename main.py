import csv
import os
import locale
from tkinter import *
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from datetime import datetime

try:
    import winsound
    def beep_sucesso(janela=None): winsound.MessageBeep(winsound.MB_OK)
    def beep_erro(janela=None): winsound.MessageBeep(winsound.MB_ICONHAND)
except ImportError:
    def beep_sucesso(janela=None):
        if janela: janela.bell()
    def beep_erro(janela=None):
        if janela: janela.bell()

# Locale fix for platforms
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'pt_BR')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')
        except locale.Error:
            locale.setlocale(locale.LC_ALL, '')

ARQUIVO_ESTOQUE = "estoque.csv"
ARQUIVO_HISTORICO = "historico_vendas.csv"

def atualizar_rodape(label, estoque):
    total = len(estoque)
    if total == 0:
        media_preco = 0
        soma_preco = 0
        media_ano = 0
    else:
        soma_preco = sum(c['Preço'] for c in estoque)
        media_preco = soma_preco / total
        media_ano = sum(c['Ano'] for c in estoque) / total
    texto = (
        f"Total de veículos: {total}    |    "
        f"Média de preço: {locale.currency(media_preco, grouping=True)}    |    "
        f"Valor total: {locale.currency(soma_preco, grouping=True)}    |    "
        f"Média de ano: {media_ano:.0f}"
    )
    label.config(text=texto)

def mostrar_feedback(label, mensagem, cor="green", tempo=2000):
    label.config(text=mensagem, fg=cor)
    label.after(tempo, lambda: label.config(text=""))

def centralizar_janela(janela, largura, altura):
    janela.update_idletasks()
    largura_tela = janela.winfo_screenwidth()
    altura_tela = janela.winfo_screenheight()
    x = (largura_tela // 2) - (largura // 2)
    y = (altura_tela // 2) - (altura // 2)
    janela.geometry(f"{largura}x{altura}+{x}+{y}")

def abrir_arquivo_default(path):
    import sys
    import subprocess
    try:
        if os.name == 'nt':
            os.startfile(path)
        elif sys.platform == 'darwin':
            subprocess.run(['open', path])
        else:
            subprocess.run(['xdg-open', path])
    except Exception as e:
        messagebox.showwarning("Aviso", f"Não foi possível abrir o arquivo automaticamente.\nAbra manualmente: {path}")

def carregar_estoque():
    estoque = []
    if not os.path.exists(ARQUIVO_ESTOQUE):
        with open(ARQUIVO_ESTOQUE, "w", newline='', encoding="utf-8") as f:
            writer = csv.writer(f, delimiter=';')
            writer.writerow(["Modelo", "Ano", "Preço"])
        return estoque
    with open(ARQUIVO_ESTOQUE, "r", newline='', encoding="utf-8") as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            try:
                if row['Modelo']:
                    estoque.append({
                        "Modelo": row["Modelo"],
                        "Ano": int(row["Ano"]),
                        "Preço": float(row["Preço"])
                    })
            except (ValueError, KeyError):
                continue
    return estoque

def salvar_estoque(estoque):
    with open(ARQUIVO_ESTOQUE, "w", newline='', encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=';')
        writer.writerow(["Modelo", "Ano", "Preço"])
        for carro in estoque:
            writer.writerow([carro["Modelo"], carro["Ano"], carro["Preço"]])

def salvar_historico_venda(carro):
    novo_arquivo = not os.path.exists(ARQUIVO_HISTORICO)
    with open(ARQUIVO_HISTORICO, "a", newline='', encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=';')
        if novo_arquivo:
            writer.writerow(["Modelo", "Ano", "Preço", "Data"])
        datahora = datetime.now().strftime("%d/%m/%Y")
        writer.writerow([carro["Modelo"], carro["Ano"], carro["Preço"], datahora])

def atualizar_treeview(tree, estoque, rodape=None):
    tree.delete(*tree.get_children())
    for idx, carro in enumerate(estoque):
        preco = locale.currency(carro["Preço"], grouping=True)
        tree.insert("", "end", iid=idx, text=idx, values=(carro["Modelo"], carro["Ano"], preco))
    if rodape is not None:
        atualizar_rodape(rodape, estoque)

def adicionar_carro_gui(tree, estoque, rodape, feedback_label, root):
    def salvar():
        try:
            modelo = entry_modelo.get().strip()
            ano = entry_ano.get().strip()
            preco = entry_preco.get().strip().replace(",", ".")
            if not modelo:
                mostrar_feedback(feedback_label, "Modelo não pode ser vazio.", "red")
                beep_erro(win_add)
                return
            try:
                ano = int(ano)
                if ano < 1900 or ano > 2100:
                    raise ValueError
            except ValueError:
                mostrar_feedback(feedback_label, "Ano inválido.", "red")
                beep_erro(win_add)
                return
            try:
                preco = float(preco)
                if preco < 0:
                    raise ValueError
            except ValueError:
                mostrar_feedback(feedback_label, "Preço inválido.", "red")
                beep_erro(win_add)
                return
            estoque.append({"Modelo": modelo, "Ano": ano, "Preço": preco})
            salvar_estoque(estoque)
            atualizar_treeview(tree, estoque, rodape)
            mostrar_feedback(feedback_label, f"Carro '{modelo}' adicionado com sucesso!", "green")
            beep_sucesso(win_add)
            win_add.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro inesperado: {e}")
            win_add.destroy()
    win_add = Toplevel()
    win_add.title("Adicionar Carro")
    centralizar_janela(win_add, 350, 200)
    win_add.resizable(False, False)
    win_add.grab_set()
    Label(win_add, text="Modelo:", font=('Arial', 13)).grid(row=0, column=0, sticky="e", padx=8, pady=8)
    entry_modelo = Entry(win_add, font=('Arial', 13))
    entry_modelo.grid(row=0, column=1, padx=8, pady=8)
    Label(win_add, text="Ano:", font=('Arial', 13)).grid(row=1, column=0, sticky="e", padx=8, pady=8)
    entry_ano = Entry(win_add, font=('Arial', 13))
    entry_ano.grid(row=1, column=1, padx=8, pady=8)
    Label(win_add, text="Preço:", font=('Arial', 13)).grid(row=2, column=0, sticky="e", padx=8, pady=8)
    entry_preco = Entry(win_add, font=('Arial', 13))
    entry_preco.grid(row=2, column=1, padx=8, pady=8)
    Button(win_add, text="Salvar", font=('Arial', 12), width=13, command=salvar).grid(row=3, column=0, columnspan=2, pady=14)

def editar_carro_gui(tree, estoque, rodape, feedback_label, root):
    selected = tree.focus()
    if not selected:
        mostrar_feedback(feedback_label, "Selecione um carro para editar.", "red")
        beep_erro(root)
        return
    idx = int(selected)
    carro = estoque[idx]
    def salvar():
        try:
            modelo = entry_modelo.get().strip()
            ano = entry_ano.get().strip()
            preco = entry_preco.get().strip().replace(",", ".")
            if not modelo:
                mostrar_feedback(feedback_label, "Modelo não pode ser vazio.", "red")
                beep_erro(win_edit)
                return
            try:
                ano = int(ano)
                if ano < 1900 or ano > 2100:
                    raise ValueError
            except ValueError:
                mostrar_feedback(feedback_label, "Ano inválido.", "red")
                beep_erro(win_edit)
                return
            try:
                preco = float(preco)
                if preco < 0:
                    raise ValueError
            except ValueError:
                mostrar_feedback(feedback_label, "Preço inválido.", "red")
                beep_erro(win_edit)
                return
            carro["Modelo"] = modelo
            carro["Ano"] = ano
            carro["Preço"] = preco
            salvar_estoque(estoque)
            atualizar_treeview(tree, estoque, rodape)
            mostrar_feedback(feedback_label, "Carro editado com sucesso!", "green")
            beep_sucesso(win_edit)
            win_edit.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro inesperado: {e}")
            win_edit.destroy()
    win_edit = Toplevel()
    win_edit.title("Editar Carro")
    centralizar_janela(win_edit, 350, 200)
    win_edit.resizable(False, False)
    win_edit.grab_set()
    Label(win_edit, text="Modelo:", font=('Arial', 13)).grid(row=0, column=0, sticky="e", padx=8, pady=8)
    entry_modelo = Entry(win_edit, font=('Arial', 13))
    entry_modelo.insert(0, carro["Modelo"])
    entry_modelo.grid(row=0, column=1, padx=8, pady=8)
    Label(win_edit, text="Ano:", font=('Arial', 13)).grid(row=1, column=0, sticky="e", padx=8, pady=8)
    entry_ano = Entry(win_edit, font=('Arial', 13))
    entry_ano.insert(0, carro["Ano"])
    entry_ano.grid(row=1, column=1, padx=8, pady=8)
    Label(win_edit, text="Preço:", font=('Arial', 13)).grid(row=2, column=0, sticky="e", padx=8, pady=8)
    entry_preco = Entry(win_edit, font=('Arial', 13))
    entry_preco.insert(0, carro["Preço"])
    entry_preco.grid(row=2, column=1, padx=8, pady=8)
    Button(win_edit, text="Salvar", font=('Arial', 12), width=13, command=salvar).grid(row=3, column=0, columnspan=2, pady=14)

def excluir_carro_gui(tree, estoque, rodape, feedback_label, root):
    selected = tree.focus()
    if not selected:
        mostrar_feedback(feedback_label, "Selecione um carro para excluir.", "red")
        beep_erro(root)
        return
    idx = int(selected)
    carro = estoque[idx]
    confirm = messagebox.askyesno("Confirmar Exclusão", f"Excluir '{carro['Modelo']} {carro['Ano']}'?")
    if confirm:
        salvar_historico_venda(carro)
        estoque.pop(idx)
        salvar_estoque(estoque)
        atualizar_treeview(tree, estoque, rodape)
        mostrar_feedback(feedback_label, "Carro excluído.", "green")
        beep_sucesso(root)

def gerar_arquivo_para_impressao(estoque):
    if not estoque:
        messagebox.showerror("Erro", "Estoque vazio! Nenhum arquivo gerado.")
        return
    doc = Document()
    doc.add_heading("Estoque de Veículos", 0)
    tabela = doc.add_table(rows=1, cols=6)
    tabela.style = "Table Grid"
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Modelo'
    hdr_cells[1].text = 'Ano'
    hdr_cells[2].text = 'Preço (R$)'
    hdr_cells[3].text = 'Fotos'
    hdr_cells[4].text = 'Site'
    hdr_cells[5].text = 'Stories'
    col_widths = [Cm(7.7), Cm(2.5), Cm(3.8), Cm(0.3), Cm(0.3), Cm(0.3)]
    for i, width in enumerate(col_widths):
        hdr_cells[i].width = width
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = Pt(16)
    for carro in estoque:
        row_cells = tabela.add_row().cells
        row_cells[0].text = carro["Modelo"]
        row_cells[1].text = str(carro["Ano"])
        row_cells[2].text = locale.currency(carro["Preço"], grouping=True)
        row_cells[3].text = ""
        row_cells[4].text = ""
        row_cells[5].text = ""
        for cell in row_cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = Pt(20)
    doc.add_paragraph("\n")
    nome_arquivo = "estoque_para_impressao.docx"
    doc.save(nome_arquivo)
    messagebox.showinfo("Sucesso", f"Arquivo '{nome_arquivo}' gerado com sucesso!")
    beep_sucesso()
    abrir_arquivo_default(nome_arquivo)

def visualizar_historico_vendas():
    if not os.path.exists(ARQUIVO_HISTORICO):
        messagebox.showinfo("Histórico de Vendas", "Nenhuma venda registrada ainda.")
        return
    abrir_arquivo_default(ARQUIVO_HISTORICO)

def main():
    estoque = carregar_estoque()
    root = Tk()
    root.title("Estoque de Veículos - Loja")
    root.geometry("750x500")
    style = ttk.Style(root)
    style.configure("Treeview.Heading", font=('Arial', 15, 'bold'))
    style.configure("Treeview", font=('Arial', 14), rowheight=38, foreground='black')
    style.layout("Treeview", [('Treeview.treearea', {'sticky': 'nswe'})])
    style.map('Treeview', background=[('selected', '#ececec')])
    Label(root, text="Controle de Estoque de Veículos", font=("Arial", 19, "bold")).pack(pady=10)
    frame = Frame(root)
    frame.pack(fill=BOTH, expand=True)
    columns = ("Modelo", "Ano", "Preço")
    tree = ttk.Treeview(frame, columns=columns, show="headings")
    sort_state = {"Ano": True, "Preço": True}
    feedback_label = Label(root, font=('Arial', 12, "bold"), fg="green", anchor="center", justify="center")
    feedback_label.pack(fill="x", pady=(8, 2))
    rodape = Label(root, font=('Arial', 13, "bold"), anchor="center", justify="center", pady=10)
    rodape.pack(fill="x")
    def sort_column(col):
        reverse = not sort_state[col]
        sort_state[col] = reverse
        if col == "Ano":
            estoque.sort(key=lambda x: x['Ano'], reverse=reverse)
        elif col == "Preço":
            estoque.sort(key=lambda x: x['Preço'], reverse=reverse)
        atualizar_treeview(tree, estoque, rodape)
        mostrar_feedback(feedback_label, f"Ordenado por {col} {'↓' if reverse else '↑'}", "blue")
    tree.heading("Modelo", text="Modelo", anchor="center")
    tree.heading("Ano", text="Ano", anchor="center", command=lambda: sort_column("Ano"))
    tree.heading("Preço", text="Preço", anchor="center", command=lambda: sort_column("Preço"))
    tree.column("Modelo", width=320, anchor="center")
    tree.column("Ano", width=100, anchor="center")
    tree.column("Preço", width=170, anchor="center")
    tree.pack(side=LEFT, fill=BOTH, expand=True)
    scrollbar = ttk.Scrollbar(frame, orient=VERTICAL, command=tree.yview)
    tree.configure(yscroll=scrollbar.set)
    scrollbar.pack(side=RIGHT, fill=Y)
    atualizar_treeview(tree, estoque, rodape)
    btn_frame = Frame(root)
    btn_frame.pack(pady=12)
    Button(btn_frame, text="Adicionar Carro", font=('Arial', 13), width=18,
           command=lambda: adicionar_carro_gui(tree, estoque, rodape, feedback_label, root)).grid(row=0, column=0, padx=8)
    Button(btn_frame, text="Editar Carro", font=('Arial', 13), width=18,
           command=lambda: editar_carro_gui(tree, estoque, rodape, feedback_label, root)).grid(row=0, column=1, padx=8)
    Button(btn_frame, text="Excluir Carro", font=('Arial', 13), width=18,
           command=lambda: excluir_carro_gui(tree, estoque, rodape, feedback_label, root)).grid(row=0, column=2, padx=8)
    Button(btn_frame, text="Gerar Arquivo para Impressão", font=('Arial', 13), width=24,
           command=lambda: gerar_arquivo_para_impressao(estoque)).grid(row=0, column=3, padx=8)
    Button(btn_frame, text="Histórico de Vendas", font=('Arial', 13), width=18,
           command=visualizar_historico_vendas).grid(row=0, column=4, padx=8)
    Button(btn_frame, text="Sair", font=('Arial', 13), width=10, command=root.destroy).grid(row=0, column=5, padx=8)
    root.mainloop()

if __name__ == "__main__":
    main()